"""Export transcript to TXT, DOCX, PDF, SRT, VTT."""
from datetime import datetime
import io
from typing import Tuple, Optional, List, Dict, Any

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY


def format_duration(seconds: float) -> str:
    """Format duration in seconds to human-readable string."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    if hours > 0:
        return f"{hours} giờ {minutes} phút {secs} giây"
    if minutes > 0:
        return f"{minutes} phút {secs} giây"
    return f"{secs} giây"


def export_txt(transcript: str, filename: str = "transcript.txt") -> Tuple[bytes, str]:
    """Export transcript to TXT. Returns (bytes, filename)."""
    return transcript.encode("utf-8"), filename


def export_docx(
    transcript: str,
    metadata: Optional[dict] = None,
    filename: str = "transcript.docx",
) -> Tuple[bytes, str]:
    """Export transcript to DOCX. Returns (bytes, filename)."""
    doc = Document()
    doc.add_heading("Bản Ghi Âm Thanh", 0)
    if metadata:
        doc.add_paragraph(f"Thời gian: {metadata.get('timestamp', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}")
        if "duration" in metadata:
            doc.add_paragraph(f"Độ dài: {format_duration(metadata['duration'])}")
        if "word_count" in metadata:
            doc.add_paragraph(f"Số từ: {metadata['word_count']}")
        doc.add_paragraph("")
    for para in transcript.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue(), filename


def export_pdf(
    transcript: str,
    metadata: Optional[dict] = None,
    filename: str = "transcript.pdf",
) -> Tuple[bytes, str]:
    """Export transcript to PDF. Returns (bytes, filename)."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        rightMargin=72, leftMargin=72,
        topMargin=72, bottomMargin=18,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Heading1"],
        fontSize=16,
        textColor="#1f4e79",
        spaceAfter=12,
        alignment=TA_LEFT,
    )
    normal_style = ParagraphStyle(
        "CustomNormal",
        parent=styles["Normal"],
        fontSize=11,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=6,
    )
    story = [Paragraph("Bản Ghi Âm Thanh", title_style), Spacer(1, 12)]
    if metadata:
        meta_text = f"<b>Thời gian:</b> {metadata.get('timestamp', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}<br/>"
        if "duration" in metadata:
            meta_text += f"<b>Độ dài:</b> {format_duration(metadata['duration'])}<br/>"
        if "word_count" in metadata:
            meta_text += f"<b>Số từ:</b> {metadata['word_count']}<br/>"
        story.append(Paragraph(meta_text, normal_style))
        story.append(Spacer(1, 12))
    for para in transcript.split("\n"):
        if para.strip():
            escaped = para.strip().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(escaped, normal_style))
            story.append(Spacer(1, 6))
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue(), filename


def _seconds_to_srt_time(seconds: float) -> str:
    """Format seconds to SRT time 00:00:00,000."""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds % 1) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def _seconds_to_vtt_time(seconds: float) -> str:
    """Format seconds to VTT time 00:00:00.000."""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds % 1) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d}.{ms:03d}"


def export_srt(
    segments: List[Dict[str, Any]],
    filename: str = "subtitles.srt",
    dual: bool = False,
) -> Tuple[bytes, str]:
    """
    Export segments to SRT. Each segment: {start, end, text, optional translated_text}.
    If dual=True, each subtitle shows two lines: source then translation.
    """
    lines = []
    for i, seg in enumerate(segments, 1):
        start = seg.get("start", 0)
        end = seg.get("end", 0)
        text = (seg.get("text") or "").strip()
        trans = (seg.get("translated_text") or "").strip()
        lines.append(str(i))
        lines.append(f"{_seconds_to_srt_time(start)} --> {_seconds_to_srt_time(end)}")
        if dual and trans:
            lines.append(text)
            lines.append(trans)
        else:
            lines.append(text or trans)
        lines.append("")
    return "\n".join(lines).encode("utf-8"), filename


def export_vtt(
    segments: List[Dict[str, Any]],
    filename: str = "subtitles.vtt",
    dual: bool = False,
) -> Tuple[bytes, str]:
    """
    Export segments to WebVTT. Each segment: {start, end, text, optional translated_text}.
    If dual=True, each cue has two lines: source then translation.
    """
    lines = ["WEBVTT", ""]
    for seg in segments:
        start = seg.get("start", 0)
        end = seg.get("end", 0)
        text = (seg.get("text") or "").strip()
        trans = (seg.get("translated_text") or "").strip()
        lines.append(f"{_seconds_to_vtt_time(start)} --> {_seconds_to_vtt_time(end)}")
        if dual and trans:
            lines.append(text)
            lines.append(trans)
        else:
            lines.append(text or trans)
        lines.append("")
    return "\n".join(lines).encode("utf-8"), filename
